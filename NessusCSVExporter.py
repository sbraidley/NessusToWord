#!/usr/bin/python3
#
# Wednesday 16th March 2022
# Version: 1.2.0
# Python script to extract required Nessus data for Cyber Essentials reporting. Automating the majority of the process, however manual adjustments are required due limitations of Python-Docx.
#
# Requires python-docx and that Nessus CVS files are outputted with both CVSSv2 and CVSSv3 scores.
#
# Sam Braidley (sam.braidley@orangecyberdefense.com | samb@sensepost.com)
#
# ---------------------------------------------------------------------------------------------------------------------------------------------------
#
# Default Structure of Nessus CSV File
# Rows [Plugin ID, CVE, CVSS, Risk, Host, Protocol, Port, Name, Synopsis, Description, Solution, See Also, Plugin Output, STIG Severity, CVSS v3.0 Base Score, CVSS v2.0 Temporal Score, CVSS v3.0 Temporal Score, Risk Factor, BID, XREF, MSKB, Plugin Publication Date, Plugin Modification Date, Metasploit, Core Impact, CANVAS]

#
# To-Do: 
#
# Known Issues/Limitations:
# - Title must be formatted manually using Style2 (Limitation of Python-DocX)

import sys
import csv
import os
import docx
import operator
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

document = Document()
useMerged = False

try:
	sys.argv[1]
except IndexError:
	print('Invalid parameters, include the CVS you wish to parse ./NessusCSVExporter.py input.csv R00 or use -merge to merge multiple CSV files from Nessus.')
	sys.exit(2)
	
if "-h" in sys.argv[1]:
	print("NessusCSVExporter.py")
	print("")
	print("Uses CSV files exported from Tenable Nessus and converts them into a Microsoft DocX file for use within custom reporting templates.")
	print("")
	print("-merge : Merges all CSV files in the directory of the script and uses merged file as input.")
	sys.exit(2)

if "-merge" in sys.argv[1]:
	os.system("for filename in $(ls *.csv); do sed 1d $filename >> temp.csv; done")
	
	with open('temp.csv') as infile, open('merged.csv', 'w') as outfile:
		outfile = csv.writer(outfile)
		reader = csv.reader(infile)
		
		outfile.writerow(["Plugin ID", "CVE", "CVSS v2.0 Base Score", "Risk", "Host", "Protocol", "Port", "Name", "Synopsis", "Description", "Solution", "See Also", "Plugin Output", "STIG Severity", "CVSS v3.0 Base Score", "CVSS v2.0 Temporal Score", "CVSS v3.0 Temporal Score", "VPR Score", "Risk Factor", "References", "Plugin Information", "Exploitable With"])
		
		for row in reader:
			outfile.writerow(row)
			
	os.remove('temp.csv')
	csv_file = 'merged.csv'
	useMerged = True
elif ".csv" in sys.argv[1]:
	csv_file = sys.argv[1]
else:
	print("Invalid parameters, please ensure that the file is a CSV file ./NessusCSVExporter.py input.csv R00")
	sys.exit(2)

try:
	sys.argv[2]
	if "R" in sys.argv[2]:
		Temp_R = sys.argv[2]
		R_Number = Temp_R[1:]
	elif "R" in sys.argv[3]:
		Temp_R = sys.argv[3]
		R_Number = Temp_R[1:]
except IndexError:
	print("No reference start number provided, assuming R01.")
	R_Number = 1

# Removes large spaces
with open(csv_file) as infile, open('temp.csv', 'w') as outfile:
    outfile = csv.writer(outfile)
    for row in csv.reader(infile):
        outfile.writerow([c.replace('     ', ' ') for c in row])
		
# Remove double spaces
with open('temp.csv') as infile, open('tempa.csv', 'w') as outfile:
    outfile = csv.writer(outfile)
    for row in csv.reader(infile):
        outfile.writerow([c.replace('  ', ' ') for c in row])
		
# Remove space prior to colons
with open('tempa.csv') as infile, open('tempb.csv', 'w') as outfile:
	outfile = csv.writer(outfile)
	for row in csv.reader(infile):
		outfile.writerow([c.replace(' :', ':') for c in row])		

# Check if CVSSv3 score exists, if not use CVSSv2 score - CVSSv2 score is copied to CVSSv3 row as to not break further processing
with open('temp0.csv', 'w') as output, open ('tempb.csv') as inputfile:
	outfile = csv.writer(output)
	reader = csv.reader(inputfile)
	header = next(reader)
	
	outfile.writerow(header)
	next(reader)
	for row in reader:
		if row[14] == "":
			row[14] = row[2]
			outfile.writerow(row)
		else:
			outfile.writerow(row)
	
# Rewritten CVSS score sorting
def getCVSS(elem):
    return elem[14]

with open('temp1.csv', 'w') as output, open('temp0.csv') as inputfile:
	outfile = csv.writer(output)
	reader = csv.reader(inputfile)
	header = next(reader)

	data = []
	outfile.writerow(header)
	for row in reader:
		if row[14] == "":
			continue
		elif row[14] == "0.0":
			continue
		elif row[14] == "0":
			continue
		else:
			x = row[14]
			i = float(x)
			data.append([row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8], row[9], row[10], row[11], row[12], row[13], i, row[15], row[16], row[17], row[18], row[19], row[20], row[21]])	
	data.sort(key=getCVSS, reverse=True)
	for entry in data:
		outfile.writerow(entry)

#Builds CSV for host list merging
with open('temp1.csv') as infile, open('temp2.csv', 'w') as outfile:
    outfile = csv.writer(outfile)
    reader = csv.reader(infile)
    header = next(reader)
    
    data = []
    for row in reader:
        data.append([row[0], (row[4] + " (" + row[6]) + "/" + row[5] + ")"])
    outfile.writerow(["ID", "Host"])
    for entry in data:
        outfile.writerow(entry)
    
with open('temp1.csv') as infile, open ('temp2.csv') as hostIPList, open('temp3.csv', 'w') as outfile:
    outfile = csv.writer(outfile)
    hostIPListReader = csv.reader(hostIPList)
    reader = csv.reader(infile)
    header = next(infile)
    
    outfile.writerow(["Plugin ID", "CVE", "CVSS v2.0 Base Score", "Risk", "Host", "Protocol", "Port", "Name", "Synopsis", "Description", "Solution", "See Also", "Plugin Output", "STIG Severity", "CVSS v3.0 Base Score", "CVSS v2.0 Temporal Score", "CVSS v3.0 Temporal Score", "VPR Score", "Risk Factor", "References", "Plugin Information", "Exploitable With"])

    mylist = []
    for row in reader:
        for element in hostIPListReader:
            if row[0] == element[0]:
            	mylist.append(element[1])
            	continue
            mylist = list( dict.fromkeys(mylist) )
            row[4] = str(mylist)
        hostIPList.seek(0)
        outfile.writerow(row)
        mylist = []
		
# Temporary fix for missing hosts at end of file
lastFinding = ""
# Gets lastfinding value
with open('temp2.csv') as infile:
	for row in reversed(list(csv.reader(infile))):
		lastFinding = row[0]
		break
	
# Get all hosts for last finding value
mylist2 = []
with open('temp2.csv') as infile:
	reader = csv.reader(infile)
	for row in reader:
		if row[0] == lastFinding:
			mylist2.append(row[1])
			continue

# Use mylist2 to add to last entry in temp3a csv
with open('temp3.csv') as infile, open('temp3a.csv', 'w') as outfile:
	reader = csv.reader(infile)
	outfile = csv.writer(outfile)
	for row in reader:
		if row[0] == lastFinding:
			row[4] = str(mylist2)
			outfile.writerow(row)
		else:
			outfile.writerow(row)

# Removes duplicates
with open('temp3a.csv') as infile, open('temp4.csv', 'w') as outfile:
    outfile = csv.writer(outfile)
    seen = set()
    for row in csv.reader(infile):
        if row[0] in seen:
        	continue # skip duplicate
        seen.add(row[0])
        outfile.writerow(row)

# Generates findings table data in CSV form to be copied into Word in-line with the results generated using the tool
# Format = Reference Number | Finding Title | CVSSv3 Score
table_file = 'temp4.csv'

with open(table_file, 'rt') as csvdata, open ('tabledata.csv', 'w') as outfile:
	outfile = csv.writer(outfile)
	csvdata = csv.DictReader(csvdata)

	outfile.writerow(["Ref #", "Vulnerability", "CVSS"])

	tablecounter = int(R_Number)

	for row in csvdata:
		data = []
		if len(row['CVSS v3.0 Base Score']) < 1:
			continue
		elif float(row['CVSS v3.0 Base Score']) >= 4.0:
			if(tablecounter < 10):
				ref_no = (('R0') + str(tablecounter))
				data.append(ref_no)
			else:
				ref_no = (('R') + str(tablecounter))
				data.append(ref_no)
			finding_name = row['Name']
			data.append(finding_name)
			score = row['CVSS v3.0 Base Score']
			data.append(score)
			outfile.writerow(data)
			data = []
			tablecounter+=1
		else:
			continue

def add_hyperlink(paragraph, url, text, color):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add colour if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

csv_file = 'temp4.csv'

with open(csv_file, 'rt') as csvfile:
	
	nessusreader = csv.DictReader(csvfile)
	counter = int(R_Number)

	style = document.styles['Normal']
	font = style.font
	font.name = 'Helvetica 75 Bold'
	font.size = Pt(9)	
	font2 = style.font
	font2.name = 'Helvetica 55 Roman'
	font2.size = Pt(9)

	for row in nessusreader:
		x = row['CVSS v3.0 Base Score']
		i = float(x)
		#if(x >= 4):
		if(i >= 4):
			name = document.add_paragraph()
			nameRow = row['Name']
			if(counter<10):
				font = name.add_run('R0' + str(counter) + ' - ' + nameRow).font
			if(counter >= 10):
				font = name.add_run('R' + str(counter) + ' - ' + nameRow).font
			font.bold = True
			font.underline = True
			font.size = Pt(10)
			syn1 = row['Synopsis']
			syn2 = " ".join(syn1.splitlines())
			syn3 = syn2.rstrip("\n")
			syn4 = syn3.replace(":  ", ": \n")
			syn5 = syn4.replace("  ", " ")
			syn6 = syn5.replace("  ", " ")
			p = document.add_paragraph(syn6)
			p.style = document.styles['Normal']
			p1 = document.add_paragraph().add_run('Description').bold = True
			desc = row['Description']
			desc2 = " ".join(desc.splitlines())
			desc3 = desc2.rstrip("\n")
			desc4 = desc3.replace(":  ", ": \n")
			desc5 = desc4.replace("  ", " ")
			desc6 = desc5.replace("- ", "\n - ")
			desc7 = desc6.replace("  ", " ")
			p2 = document.add_paragraph(desc7)
			p2.style = document.styles['Normal']
			p3 = document.add_paragraph('')
			p3.add_run('Affected Hosts').bold = True
			p3.style = document.styles['Normal']
			hosts = row['Host'].split("', '")
			hosts = [s.replace('[', '') for s in hosts]
			hosts = [s.replace(']', '') for s in hosts]
			hosts = [s.replace("'", '') for s in hosts] 
			for item in hosts:
				p8 = document.add_paragraph(item, style='List Bullet')
				p8.style = document.styles['List Bullet']
			p4 = document.add_paragraph('')
			p4.add_run('Recommendations').bold = True
			rec1 = row['Solution']
			rec2 = " ".join(rec1.splitlines())
			rec3 = rec2.rstrip("\n")
			rec4 = rec3.replace(":  ", ": \n")
			rec5 = rec4.replace("  ", " ")
			rec6 = rec5.replace("-", "\n - ")
			rec7 = rec6.replace("  ", " ")
			rec8 = rec7.replace(" For", "\n\nFor")
			p5 = document.add_paragraph(rec8)
			
			p5.style = document.styles['Normal']
			if row['See Also'] == "":
				counter+=1
				continue
			elif row['See Also'] == " ":
				counter+=1
				continue
			else:
				p6 = document.add_paragraph().add_run('External References').bold = True
				links = row['See Also'].split(" ")
				for item in links:
	    				p = document.add_paragraph()
	    				run = p.add_run()
	    				font = run.font
	    				font.underline = True
	    				hyperlink = add_hyperlink(p, item, item, '0645AD')
			counter+=1

#Clean up
os.remove('temp.csv')
os.remove('tempa.csv')
os.remove('tempb.csv')
os.remove('temp1.csv')
os.remove('temp2.csv')
os.remove('temp3.csv')
os.remove('temp3a.csv')
os.remove('temp4.csv')
os.remove('temp0.csv')
if useMerged is True:
	os.remove('merged.csv')

# Save DocX
document.save('export.docx')		